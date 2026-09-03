import io
from typing import List, Dict, Any
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    """Set cell margins (padding) in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
        
    tcPr.append(tcMar)

def set_cell_borders(cell, color="1F497D", sz="12", val="single"):
    """Set border for a cell."""
    borders_xml = f'''
    <w:tcBorders {nsdecls("w")}>
        <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
    </w:tcBorders>
    '''
    cell._tc.get_or_add_tcPr().append(parse_xml(borders_xml))

def prevent_row_split(row):
    """Prevent table row from splitting across pages."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def build_page_disclaimer_header(section, usable_width):
    """
    Builds the top bilingual disclaimer header box for each page:
    Left: Tamil disclaimer
    Right: English disclaimer
    Repeats on each and every page automatically via Word section header.
    """
    header = section.header
    # Remove default empty paragraph in header
    if header.paragraphs:
        p_head = header.paragraphs[0]._element
        p_head.getparent().remove(p_head)
        
    tbl_hdr = header.add_table(rows=1, cols=2, width=usable_width)
    tbl_hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_hdr.autofit = False
    
    c_tam = tbl_hdr.cell(0, 0)
    c_eng = tbl_hdr.cell(0, 1)
    half_width = usable_width / 2
    c_tam.width = half_width
    c_eng.width = half_width
    c_tam.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    c_eng.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    set_cell_borders(c_tam, color="000000", sz="12", val="single")
    set_cell_borders(c_eng, color="000000", sz="12", val="single")
    set_cell_margins(c_tam, top=35, bottom=35, left=70, right=70)
    set_cell_margins(c_eng, top=35, bottom=35, left=70, right=70)
    
    # Left: Tamil
    p_tam = c_tam.paragraphs[0]
    p_tam.paragraph_format.space_before = Pt(0)
    p_tam.paragraph_format.space_after = Pt(0)
    p_tam.paragraph_format.line_spacing = 1.05
    r_tam = p_tam.add_run("சிகப்பு மை பேனாவினால் - இந்த PROOF PAPER-ல் மட்டும் திருத்தம் செய்யவும். கார்டு வந்த பிறகு திருத்தம் செய்தால் கண்டிப்பாக கார்டு மாற்றி தரமுடியாது.")
    r_tam.font.bold = True
    r_tam.font.size = Pt(7.5)
    r_tam.font.name = "Nirmala UI"
    rPr_tam = r_tam._r.get_or_add_rPr()
    rFonts_tam = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Nirmala UI" w:hAnsi="Nirmala UI" w:cs="Nirmala UI"/>')
    rPr_tam.append(rFonts_tam)
    rPr_tam.append(parse_xml(f'<w:bCs {nsdecls("w")}/>'))
    
    # Right: English
    p_eng = c_eng.paragraphs[0]
    p_eng.paragraph_format.space_before = Pt(0)
    p_eng.paragraph_format.space_after = Pt(0)
    p_eng.paragraph_format.line_spacing = 1.05
    r_eng = p_eng.add_run("MAKE CORRECTIONS ONLY ON THIS PROOF PAPER USING RED INK PEN. NO CORRECTIONS OR REPLACEMENTS WILL BE ENTERTAINED AFTER ID CARD IS PRINTED.")
    r_eng.font.bold = True
    r_eng.font.size = Pt(7.2)
    r_eng.font.name = "Arial"


def generate_word_document(staff_records: List[Dict[str, Any]]) -> bytes:
    """
    Generates a single Word document (Staff_ID_Proof.docx) containing all staff proof cards,
    formatted as a 2×6 grid (12 cards per A4 page: 6 rows × 2 columns), with a top bilingual
    disclaimer (Tamil left, English right) on each and every page.
    """
    doc = docx.Document()
    
    # Configure A4 Portrait setup (8.27 in x 11.69 in) with compact margins
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.80)
    section.bottom_margin = Inches(0.20)
    section.left_margin = Inches(0.40)
    section.right_margin = Inches(0.40)
    section.header_distance = Inches(0.18)
    
    # Compute card dimensions for a 6‑row × 2‑column grid (12 cards per A4 page)
    usable_width = Inches(8.27) - Inches(0.40) * 2  # 7.47 inches
    CARD_WIDTH = usable_width / 2
    CARD_HEIGHT = Inches(1.70)  # 6 rows @ 1.70 in = 10.20 in (fits comfortably in 10.69 in printable height)
    
    # Build top bilingual disclaimer on every page
    build_page_disclaimer_header(section, usable_width)
    
    # Remove default initial blank paragraph created by docx.Document()
    if doc.paragraphs:
        p0 = doc.paragraphs[0]._element
        p0.getparent().remove(p0)
    
    total_records = len(staff_records)
    if total_records == 0:
        output_stream = io.BytesIO()
        doc.save(output_stream)
        return output_stream.getvalue()
        
    # Total cards rounded up to multiple of 12 (6 rows * 2 cols = 12 cards/page)
    total_cards = ((total_records + 11) // 12) * 12
    total_rows = total_cards // 2
    
    table = doc.add_table(rows=total_rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    for r_idx in range(total_rows):
        row = table.rows[r_idx]
        prevent_row_split(row)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = CARD_HEIGHT
        
        for c_idx in range(2):
            card_idx = r_idx * 2 + c_idx
            cell = row.cells[c_idx]
            cell.width = CARD_WIDTH
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            
            if card_idx < total_records:
                record = staff_records[card_idx]
                build_staff_card(cell, record)
                set_cell_borders(cell, color="1F497D", sz="12", val="single")
                set_cell_margins(cell, top=20, bottom=20, left=50, right=50)
            else:
                # Blank placeholder cell for remaining slots on last page
                set_cell_borders(cell, color="FFFFFF", sz="0", val="none")
                
    output_stream = io.BytesIO()
    doc.save(output_stream)
    return output_stream.getvalue()


def build_staff_card(cell, record: Dict[str, Any]):
    """Populates a cell with the Staff ID Proof Card layout using dynamic record fields for 2x6 grid."""
    
    # 1. Header Bar
    p_header = cell.paragraphs[0]
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.paragraph_format.space_before = Pt(1)
    p_header.paragraph_format.space_after = Pt(1)
    
    run_header = p_header.add_run("ID PROOF")
    run_header.font.name = "Arial"
    run_header.font.size = Pt(8.5)
    run_header.font.bold = True
    run_header.font.color.rgb = RGBColor(31, 73, 125) # Navy Blue Accent
    
    # Divider rule
    p_rule = cell.add_paragraph()
    p_rule.paragraph_format.space_before = Pt(0)
    p_rule.paragraph_format.space_after = Pt(2)
    p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_rule = p_rule.add_run("─" * 38)
    r_rule.font.size = Pt(5)
    r_rule.font.color.rgb = RGBColor(190, 205, 225)

    # 2. Nested Table for Photo (Left) and Details (Right)
    nested_table = cell.add_table(rows=1, cols=2)
    nested_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    nested_table.autofit = False
    
    col_photo = nested_table.cell(0, 0)
    col_details = nested_table.cell(0, 1)
    
    col_photo.width = Inches(1.05)
    col_details.width = Inches(2.45)
    
    set_cell_margins(col_photo, top=0, bottom=0, left=5, right=10)
    set_cell_margins(col_details, top=0, bottom=0, left=10, right=5)
    
    set_cell_borders(col_photo, color="FFFFFF", sz="0", val="none")
    set_cell_borders(col_details, color="FFFFFF", sz="0", val="none")
    
    # 3. Photo Column
    p_photo = col_photo.paragraphs[0]
    p_photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_photo.paragraph_format.space_before = Pt(0)
    p_photo.paragraph_format.space_after = Pt(0)
    
    photo_bytes = record.get("_photo_bytes") or record.get("photo_bytes")
    if photo_bytes:
        photo_stream = io.BytesIO(photo_bytes)
        try:
            p_photo.add_run().add_picture(photo_stream, width=Inches(0.85))
        except Exception:
            p_photo.add_run("[Photo Error]")

    # 4. Filter & Prepare Dynamic Fields
    key_cols = record.get("_key_cols", {})
    photo_col = key_cols.get("photo_col") if isinstance(key_cols, dict) else None
    
    display_fields = []
    has_address = any(k.upper() == "ADDRESS" for k in record.keys() if not k.startswith("_"))
    
    for key, val in record.items():
        if key.startswith("_") or key.startswith("display_"):
            continue
        # Skip photo column from text list if it matches photo column name
        if photo_col and key == photo_col:
            continue
        if key.upper() in ("PHOTO", "IMAGE", "PIC", "PHOTO_NAME"):
            continue
        # Skip ADD1 / ADD2 if unified ADDRESS is present
        if has_address and key.upper() in ("ADD1", "ADD2"):
            continue
            
        str_val = str(val).strip() if val is not None else ""
        disp_val = str_val if str_val else "-"
        display_fields.append((key, disp_val))

    total_fields = len(display_fields)
    
    # Responsive font & spacing adjustment based on total fields for compact 2x6 cards
    if total_fields <= 6:
        font_sz = 7.0
        space_aft = 0.8
    elif total_fields <= 9:
        font_sz = 6.4
        space_aft = 0.4
    else:
        font_sz = 5.6
        space_aft = 0.1

    p_first = col_details.paragraphs[0]
    p_first.paragraph_format.space_before = Pt(0)
    p_first.paragraph_format.space_after = Pt(space_aft)
    p_first.paragraph_format.line_spacing = 1.02
    
    for idx, (label, value) in enumerate(display_fields):
        p_line = p_first if idx == 0 else col_details.add_paragraph()
        p_line.paragraph_format.space_before = Pt(0)
        p_line.paragraph_format.space_after = Pt(space_aft)
        p_line.paragraph_format.line_spacing = 1.02
        
        is_title = (idx == 0) or ("NAME" in label.upper() or "TITLE" in label.upper())
        is_id = ("ID" in label.upper() or "CODE" in label.upper() or "REG" in label.upper())
        
        # Label
        r_label = p_line.add_run(f"{label}: ")
        r_label.font.name = "Arial"
        r_label.font.size = Pt(font_sz)
        r_label.font.bold = True
        r_label.font.color.rgb = RGBColor(50, 50, 50)
        
        # Value
        r_val = p_line.add_run(str(value))
        r_val.font.name = "Arial"
        r_val.font.size = Pt(font_sz)
        
        if is_title and idx < 2:
            r_val.font.bold = True
            r_val.font.color.rgb = RGBColor(31, 73, 125) # Navy Highlight for primary name/title
        elif is_id and idx < 3:
            r_val.font.bold = True
            r_val.font.color.rgb = RGBColor(0, 0, 0)
        else:
            r_val.font.bold = False
            r_val.font.color.rgb = RGBColor(30, 30, 30)

    # 5. "No of corrections _____" line at bottom right of each card
    p_corr = cell.paragraphs[-1] if len(cell.paragraphs) > 2 else cell.add_paragraph()
    p_corr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_corr.paragraph_format.space_before = Pt(1)
    p_corr.paragraph_format.space_after = Pt(0)
    p_corr.paragraph_format.line_spacing = 1.0
    r_corr = p_corr.add_run("No of corrections _____")
    r_corr.font.name = "Arial"
    r_corr.font.size = Pt(6.8)
    r_corr.font.bold = True
    r_corr.font.color.rgb = RGBColor(60, 60, 60)

